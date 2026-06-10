import io
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

# MIME type check for binary file validation
ALLOWED_MIME_TYPES = {
    ".pdf": b"%PDF",
    ".docx": b"PK\x03\x04",
}


class ParseError(Exception):
    """Raised when document parsing fails."""
    pass


class UnsupportedFileTypeError(ParseError):
    """Raised when the file type is not supported."""
    pass


class FileTooLargeError(ParseError):
    """Raised when the file exceeds the size limit."""
    pass


class FileContentInvalidError(ParseError):
    """Raised when the file content doesn't match its claimed extension."""
    pass


def validate_file(filename: str, content: bytes) -> Path:
    """Validate file type, size, and content integrity. Returns the file path object."""
    if not filename:
        raise UnsupportedFileTypeError("文件名为空")

    ext = Path(filename).suffix.lower()

    if not ext:
        raise UnsupportedFileTypeError("无法识别文件类型，缺少文件扩展名")

    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(
            f"不支持的文件类型: {ext}。支持的类型: {supported}"
        )

    if len(content) == 0:
        raise FileContentInvalidError("文件内容为空，请上传有效文件")

    if len(content) > MAX_FILE_SIZE_BYTES:
        max_mb = MAX_FILE_SIZE_BYTES // (1024 * 1024)
        raise FileTooLargeError(
            f"文件大小 ({len(content) / (1024*1024):.1f}MB) 超过限制 ({max_mb}MB)"
        )

    # Validate binary file magic bytes
    expected_magic = ALLOWED_MIME_TYPES.get(ext)
    if expected_magic and not content.startswith(expected_magic):
        raise FileContentInvalidError(
            f"文件内容与扩展名 {ext} 不匹配，文件可能已损坏或类型不正确"
        )

    return Path(filename)


# ---------- Parsers ----------

def parse_pdf(content: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        if len(pdf.pages) == 0:
            raise ParseError("PDF 文件不包含任何页面")
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    if not text_parts:
        raise ParseError("PDF 文件中未提取到文本内容，可能是扫描件或图片型 PDF")

    return "\n\n".join(text_parts)


def parse_markdown(content: bytes) -> str:
    """Return markdown text as-is (it's already plain text)."""
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("utf-8", errors="replace")
    return text


def parse_txt(content: bytes) -> str:
    """Return plain text as-is, trying multiple encodings."""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


def parse_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(content))

    # Check if the document has any content
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text: list[str] = []

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text for cell in row.cells if cell.text.strip()]
            if row_texts:
                tables_text.append(" | ".join(row_texts))

    if not paragraphs and not tables_text:
        raise ParseError("DOCX 文件中未提取到文本内容，文件可能为空")

    result_parts: list[str] = []

    if paragraphs:
        result_parts.append("\n\n".join(paragraphs))

    if tables_text:
        result_parts.append("\n\n--- 表格内容 ---\n" + "\n".join(tables_text))

    return "\n\n".join(result_parts)


async def parse_document(filename: str, content: bytes) -> str:
    """Parse a document by filename and return extracted text.

    Raises:
        UnsupportedFileTypeError: file extension not supported
        FileTooLargeError: file exceeds size limit
        FileContentInvalidError: file content doesn't match extension
        ParseError: parsing failed or extracted empty text
    """
    file_path = validate_file(filename, content)
    ext = file_path.suffix.lower()

    parsers = {
        ".pdf": parse_pdf,
        ".md": parse_markdown,
        ".txt": parse_txt,
        ".docx": parse_docx,
    }
    parser = parsers.get(ext)
    if parser is None:
        raise UnsupportedFileTypeError(f"不支持的文件类型: {ext}")

    try:
        text = parser(content)
    except ParseError:
        raise
    except Exception as e:
        raise ParseError(f"文档解析失败 ({ext}): {str(e)}") from e

    if not text or not text.strip():
        raise ParseError("文档解析后内容为空，无法提取有效文本")

    return text
